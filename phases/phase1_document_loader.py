"""Phase 1: Document Loader - Load PDF/DOCX and text-based documents."""

import logging
from pathlib import Path
from typing import List
import pypdf
import pdfplumber
from docx import Document

from phases.phase1_models import DocumentContent

logger = logging.getLogger(__name__)


# Text-based formats read straight off disk. API specs commonly arrive as
# .json/.yaml (OpenAPI, AsyncAPI) or .graphql/.gql (GraphQL SDL) or .proto
# (Protocol Buffers), and docs as .md/.html, so all of them route through
# load_txt — none of these formats gets special parsing, just plain text.
TEXT_SUFFIXES = {
    ".txt", ".json", ".yaml", ".yml", ".md", ".markdown", ".html", ".htm",
    ".csv", ".xml", ".graphql", ".gql", ".proto",
}


def load_document(file_path: str) -> DocumentContent:
    """Load a document and extract its content.

    Handles PDF, DOCX, and any text-based format (see TEXT_SUFFIXES). Files with
    no extension or an unrecognised one are attempted as text rather than
    rejected, since a downloaded URL often has no meaningful suffix.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(str(file_path))
    elif suffix == ".docx":
        return load_docx(str(file_path))
    elif suffix in TEXT_SUFFIXES:
        return load_txt(str(file_path))
    else:
        # Unknown suffix: try reading as text before giving up.
        logger.info(f"Unrecognised suffix '{suffix}' — attempting to read as text")
        try:
            return load_txt(str(file_path))
        except (IOError, UnicodeDecodeError) as e:
            raise ValueError(
                f"Unsupported file format: {suffix or '(none)'}. "
                f"Supported: .pdf, .docx, and text formats ({', '.join(sorted(TEXT_SUFFIXES))})"
            ) from e


def load_pdf(file_path: str) -> DocumentContent:
    """Load PDF file and extract text, tables, and metadata."""
    logger.info(f"Loading PDF: {file_path}")

    file_path = Path(file_path)
    raw_text = ""
    tables = []
    metadata = {}

    token_count = 0

    # Extract text with pypdf
    try:
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            metadata = reader.metadata or {}

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                raw_text += f"--- Page {page_num + 1} ---\n{page_text}\n"
                token_count += len(page_text.split())

    except Exception as e:
        logger.error(f"Error reading PDF with pypdf: {e}")
        raise IOError(f"Failed to read PDF: {file_path}") from e

    # Extract tables with pdfplumber
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # pdfplumber exposes extract_tables(); there is no `.tables`
                # attribute — reading one raised AttributeError, which the
                # except below swallowed, so every PDF table was silently lost.
                for table in page.extract_tables() or []:
                    tables.append(table)
                    logger.debug(f"Extracted table from page {page_num + 1}")

    except Exception as e:
        logger.warning(f"Could not extract tables from PDF: {e}")

    logger.info(f"Extracted {len(raw_text)} characters, {len(tables)} tables")

    return DocumentContent(
        file_name=file_path.name,
        raw_text=raw_text,
        tables=tables,
        metadata=dict(metadata),
        estimated_tokens=token_count,
    )


def load_docx(file_path: str) -> DocumentContent:
    """Load DOCX file and extract text and tables."""
    logger.info(f"Loading DOCX: {file_path}")

    file_path = Path(file_path)
    raw_text = ""
    tables = []
    metadata = {}

    try:
        doc = Document(file_path)

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                raw_text += para.text + "\n"

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # Extract metadata (if available)
        if doc.core_properties:
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
            }

        token_count = len(raw_text.split())

    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise IOError(f"Failed to read DOCX: {file_path}") from e

    logger.info(f"Extracted {len(raw_text)} characters, {len(tables)} tables")

    return DocumentContent(
        file_name=file_path.name,
        raw_text=raw_text,
        tables=tables,
        metadata=metadata,
        estimated_tokens=token_count,
    )


def load_txt(file_path: str) -> DocumentContent:
    """Load TXT file and extract text."""
    logger.info(f"Loading TXT: {file_path}")

    file_path = Path(file_path)

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            raw_text = f.read()

        token_count = len(raw_text.split())

    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                raw_text = f.read()
            token_count = len(raw_text.split())
        except Exception as e:
            logger.error(f"Error reading TXT with different encodings: {e}")
            raise IOError(f"Failed to read TXT: {file_path}") from e

    except Exception as e:
        logger.error(f"Error reading TXT: {e}")
        raise IOError(f"Failed to read TXT: {file_path}") from e

    logger.info(f"Extracted {len(raw_text)} characters")

    return DocumentContent(
        file_name=file_path.name,
        raw_text=raw_text,
        tables=[],
        metadata={},
        estimated_tokens=token_count,
    )


def validate_file_type(file_path: str) -> bool:
    """Check if the file is a supported type.

    Kept in step with `load_document`: PDF, DOCX, and every text format in
    TEXT_SUFFIXES. An extension this rejects is one `load_document` would only
    attempt as a last resort.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    return suffix == ".pdf" or suffix == ".docx" or suffix in TEXT_SUFFIXES


def load_multiple_documents(file_paths: List[str], llm_client=None) -> List[DocumentContent]:
    """Load multiple document files.

    When `llm_client` is supplied, embedded images are extracted and captioned,
    and the captions are appended to the document text so image content becomes
    retrievable. Without a client, images are skipped and a warning is logged —
    silently dropping them is what this pipeline is trying to stop doing.
    """
    documents = []

    for file_path in file_paths:
        try:
            doc = load_document(file_path)
            _attach_image_captions(doc, file_path, llm_client)
            documents.append(doc)
            logger.info(f"✓ Loaded: {file_path}")
        except Exception as e:
            logger.error(f"✗ Failed to load {file_path}: {e}")
            raise

    logger.info(f"Loaded {len(documents)} documents successfully")
    return documents


def _attach_image_captions(doc: DocumentContent, file_path: str, llm_client) -> None:
    """Extract, caption, and splice image content into the document text.

    Records how many images were found and how many produced a caption, on every
    path including the ones that give up early. Captioning degrades to a warning
    rather than an error, so without these counts a client with no image support
    loses every diagram while the run still reports success — the exact silent
    gap image support was built to close.
    """
    from phases.phase1_image_extractor import (
        captions_as_text,
        caption_images,
        extract_images_from_docx,
        extract_images_from_pdf,
    )

    doc.metadata.setdefault("images_found", 0)
    doc.metadata.setdefault("images_captioned", 0)

    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        images = extract_images_from_pdf(file_path, doc.file_name)
    elif suffix == ".docx":
        images = extract_images_from_docx(file_path, doc.file_name)
    else:
        return  # no embedded images in text formats

    doc.metadata["images_found"] = len(images)

    if not images:
        return

    if llm_client is None:
        logger.warning(
            f"{doc.file_name}: {len(images)} image(s) found but no LLM client was "
            f"supplied — their content will not appear in the ontology"
        )
        return

    caption_images(images, llm_client)
    doc.images = [i.data for i in images]
    # Counted from the captions themselves rather than from what captioning
    # reported, so a client without describe_image() still counts as zero.
    doc.metadata["images_captioned"] = sum(1 for i in images if i.caption)

    text = captions_as_text(images)
    if text:
        doc.raw_text = f"{doc.raw_text}\n\n{text}"
        doc.estimated_tokens = len(doc.raw_text.split())
        logger.info(f"{doc.file_name}: added {len(text)} chars of image captions")
